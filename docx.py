# LER ARQUIVO DOCX

# pip install python-docx
import docx
#https://www.bing.com/videos/search?q=extract+table+from+docx+python&&view=detail&mid=7223ED918F71E0EA3A127223ED918F71E0EA3A12&&FORM=VRDGAR&ru=%2Fvideos%2Fsearch%3Fq%3Dextract%2520table%2520from%2520docx%2520python%26qs%3Dn%26form%3DQBVRMH%26%3D%2525eManage%2520Your%2520Search%2520History%2525E%26sp%3D-1%26ghc%3D1%26pq%3Dextract%2520table%2520from%2520docx%2520python%26sc%3D1-30%26sk%3D%26cvid%3D93AFE99FB80B450396DC07872CD5485F
from docx import Document
document = Document(r'C:\pastas')
table = document.tables[0]
print(table)

lista = []
for table in document.tables:
    #print(table)
    for rows in table.rows:
        #print(rows)
        for cells in rows.cells:
            a = cells.text
            if a not in lista:
                lista.append(a)
           #celulas = cells.text

print(lista)